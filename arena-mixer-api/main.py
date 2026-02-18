from fastapi import FastAPI, UploadFile, File, Form
from fastapi.responses import StreamingResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import random
import io
import re
import traceback
import zipfile
import json
import csv
import copy
from openpyxl import Workbook

app = FastAPI(title="Arena Mix - Final Layout Engine")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], 
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

WORD_NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

# =====================================================================
# MODULE 1: CORE UTILS & BOLDING ENGINE
# =====================================================================

def get_text_from_element(element):
    return "".join(node.text for node in element.iter() if node.tag.endswith('t') and node.text)

def make_run_bold(r):
    rPr = r.find(f'{{{WORD_NS["w"]}}}rPr')
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    b = rPr.find(f'{{{WORD_NS["w"]}}}b')
    if b is None:
        b = OxmlElement('w:b')
        rPr.append(b)
    bCs = rPr.find(f'{{{WORD_NS["w"]}}}bCs')
    if bCs is None:
        bCs = OxmlElement('w:bCs')
        rPr.append(bCs)

def remove_bold(r):
    rPr = r.find(f'{{{WORD_NS["w"]}}}rPr')
    if rPr is not None:
        b = rPr.find(f'{{{WORD_NS["w"]}}}b')
        if b is not None: rPr.remove(b)
        bCs = rPr.find(f'{{{WORD_NS["w"]}}}bCs')
        if bCs is not None: rPr.remove(bCs)

def check_and_clean_answer_formatting(run_element):
    is_correct = False
    rPr = run_element.find('w:rPr', namespaces=WORD_NS)
    if rPr is not None:
        color = rPr.find('w:color', namespaces=WORD_NS)
        if color is not None:
            val = color.get(f'{{{WORD_NS["w"]}}}val', '').lower()
            if val in ['ff0000', 'red', 'c00000', 'e36c09', 'e52237']:
                is_correct = True; rPr.remove(color) 
        u = rPr.find('w:u', namespaces=WORD_NS)
        if u is not None:
            val = u.get(f'{{{WORD_NS["w"]}}}val', '').lower()
            if val != 'none':
                is_correct = True; rPr.remove(u) 
    return is_correct

def clean_marker_tags(element):
    text = get_text_from_element(element)
    if re.search(r'\[P[1-4]\]', text, re.IGNORECASE):
        cleaned_text = re.sub(r'\[P[1-4]\]\s*', '', text, flags=re.IGNORECASE)
        runs = element.findall('.//w:r', namespaces=WORD_NS)
        first = True
        for run in runs:
            t_node = run.find('w:t', namespaces=WORD_NS)
            if t_node is not None:
                if first: t_node.text = cleaned_text; first = False
                else: t_node.text = ''

def create_field_code_element(field_type):
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve')
    instrText.text = f' {field_type} \\* MERGEFORMAT '
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = '1'
    fldChar3 = OxmlElement('w:fldChar'); fldChar3.set(qn('w:fldCharType'), 'end')
    return [fldChar1, instrText, fldChar2, t, fldChar3]

def analyze_complexity(el):
    has_complex = False
    for node in el.iter():
        tag = node.tag.split('}')[-1] if '}' in node.tag else node.tag
        if tag in ['object', 'oMath', 'oMathPara', 'pict', 'shape', 'drawing']:
            has_complex = True; break
    return has_complex

def clean_paragraph_for_table(p):
    for run in p.findall(f'.//{{{WORD_NS["w"]}}}r'):
        for tab in run.findall(f'.//{{{WORD_NS["w"]}}}tab'): run.remove(tab) 
    pPr = p.find(f'{{{WORD_NS["w"]}}}pPr')
    if pPr is not None:
        ind = pPr.find(f'{{{WORD_NS["w"]}}}ind')
        if ind is not None: pPr.remove(ind) 
        jc = pPr.find(f'{{{WORD_NS["w"]}}}jc')
        if jc is not None: pPr.remove(jc)

def create_invisible_table(doc, rows, cols):
    table = doc.add_table(rows=rows, cols=cols)
    tbl = table._tbl; tblPr = tbl.tblPr
    tblW = OxmlElement('w:tblW'); tblW.set(qn('w:type'), 'pct'); tblW.set(qn('w:w'), '5000'); tblPr.append(tblW)
    tblLayout = OxmlElement('w:tblLayout'); tblLayout.set(qn('w:type'), 'fixed'); tblPr.append(tblLayout)

    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is not None: tblPr.remove(tblBorders)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}'); border.set(qn('w:val'), 'none'); tblBorders.append(border)
    tblPr.append(tblBorders)
    
    pct_width = str(int(5000 / cols))
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW'); tcW.set(qn('w:type'), 'pct'); tcW.set(qn('w:w'), pct_width)
            existing_tcW = tcPr.find(qn('w:tcW'))
            if existing_tcW is not None: tcPr.remove(existing_tcW)
            tcPr.append(tcW)
    return table

# =====================================================================
# MODULE 2: AUTO-HEADER ENGINE
# =====================================================================

def build_standard_header(doc, config_data, ma_de):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    
    tblPr = table._tbl.tblPr 
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tblBorders.append(border)
    tblPr.append(tblBorders)

    for cell in table.columns[0].cells: cell.width = Cm(8.0)
    for cell in table.columns[1].cells: cell.width = Cm(9.5)
    
    c1 = table.cell(0, 0)
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_after = Pt(0)
    p1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    p1.add_run(f"SỞ GD&ĐT {config_data.get('donVi', 'LÂM ĐỒNG').upper()}\n").bold = True
    r_tr = p1.add_run(f"TRƯỜNG {config_data.get('truong', 'THCS & THPT TUY ĐỨC').upper()}\n")
    r_tr.bold = True
    p1.add_run("-----------------------\n").bold = True
    p1.add_run("Đề chính thức\n")
    
    r_pg1 = p1.add_run("(Đề thi có ")
    r_pg1.italic = True
    r_fld = p1.add_run()
    for e in create_field_code_element('NUMPAGES'): r_fld._r.append(e)
    r_pg2 = p1.add_run(" trang)")
    r_pg2.italic = True
    
    c2 = table.cell(0, 1)
    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    p2.add_run(f"KIỂM TRA {config_data.get('kyThi', 'GIỮA KÌ 1').upper()}\n").bold = True
    p2.add_run(f"MÔN THI: {config_data.get('monThi', 'TOÁN HỌC').upper()}\n").bold = True
    r_time = p2.add_run(f"Thời gian làm bài : {config_data.get('thoiGian', '90')} phút\n")
    r_time.bold = True
    r_time.italic = True
    
    box = c2.add_table(rows=1, cols=1)
    box_tblPr = box._tbl.tblPr 
    jc = OxmlElement('w:jc'); jc.set(qn('w:val'), 'center') 
    box_tblPr.append(jc)
    
    box_borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single') 
        border.set(qn('w:sz'), '4')       
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto') 
        box_borders.append(border)
    box_tblPr.append(box_borders)
    
    box.columns[0].width = Cm(4.5)
    
    bp = box.cell(0,0).paragraphs[0]
    bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bp.paragraph_format.space_before = Pt(6)
    bp.paragraph_format.space_after = Pt(6)
    r_ma = bp.add_run(f"Mã đề: {ma_de}")
    r_ma.bold = True
    r_ma.font.size = Pt(14)
    
    doc.add_paragraph() 
    p_info = doc.add_paragraph()
    p_info.add_run("Họ tên :.................................................................................. Số báo danh : ..............................................")
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_line.paragraph_format.space_after = Pt(12)
    p_line.add_run("____________________________________________________________________________________").bold = True

# =====================================================================
# MODULE 3: PARSER
# =====================================================================

def parse_docx(doc):
    body = doc._body._body
    parsed_data = {
        "header": [], "P1": [], "P1_header": [], "P2": [], "P2_header": [],
        "P3": [], "P3_header": [], "P4": [], "P4_header": []
    }
    
    current_zone = "trash" 
    current_block = []

    for element in body:
        text = get_text_from_element(element)
        text_upper = text.strip().upper()
        
        if text_upper in ["HẾT", "---HẾT---", "HẾT.", "-HẾT-", "HẾT"]:
            continue

        if "[P1]" in text_upper or re.match(r'^PHẦN\s+(I|1|MỘT)\b', text_upper):
            if current_block and current_zone in parsed_data: parsed_data[current_zone].append({'xml': current_block})
            current_zone, current_block = "P1", []; clean_marker_tags(element); parsed_data["P1_header"].append(element); continue
        elif "[P2]" in text_upper or re.match(r'^PHẦN\s+(II|2|HAI)\b', text_upper):
            if current_block and current_zone in parsed_data: parsed_data[current_zone].append({'xml': current_block})
            current_zone, current_block = "P2", []; clean_marker_tags(element); parsed_data["P2_header"].append(element); continue
        elif "[P3]" in text_upper or re.match(r'^PHẦN\s+(III|3|BA)\b', text_upper):
            if current_block and current_zone in parsed_data: parsed_data[current_zone].append({'xml': current_block})
            current_zone, current_block = "P3", []; clean_marker_tags(element); parsed_data["P3_header"].append(element); continue
        elif "[P4]" in text_upper or re.match(r'^PHẦN\s+(IV|4|BỐN)\b', text_upper):
            if current_block and current_zone in parsed_data: parsed_data[current_zone].append({'xml': current_block})
            current_zone, current_block = "P4", []; clean_marker_tags(element); parsed_data["P4_header"].append(element); continue

        if current_zone in ["P1", "P2", "P3", "P4"]:
            if re.match(r'^Câu\s+\d+[:.\s]?', text.strip(), re.IGNORECASE):
                if current_block: parsed_data[current_zone].append({'xml': current_block})
                current_block = [element]
            else:
                if current_block: current_block.append(element)

    if current_block and current_zone in ["P1", "P2", "P3", "P4"]: parsed_data[current_zone].append({'xml': current_block})
    return parsed_data

# =====================================================================
# MODULE 4: SHUFFLE & FLEXIBLE LAYOUT
# =====================================================================

def process_options_and_extract_p1_p2(doc, block, zone_type, question_text):
    pattern = r'^\s*\*?\s*([A-D])[.)]' if zone_type == "P1" else r'^\s*\*?\s*([a-d])[.)]'
    labels = ['A', 'B', 'C', 'D'] if zone_type == "P1" else ['a', 'b', 'c', 'd']
    stem, options, current_opt = [], [], None
    
    for el in block:
        if el.tag.endswith('p'):
            text = get_text_from_element(el)
            match = re.match(pattern, text, re.IGNORECASE)
            if match:
                if current_opt is not None: options.append(current_opt)
                current_opt = {'xml': [el], 'is_correct': False}
                if '*' in text or '∗' in text or re.search(r'\(\s*đ(?:úng)?\s*\)', text, re.IGNORECASE):
                    current_opt['is_correct'] = True
                
                for run in el.findall('.//w:r', namespaces=WORD_NS):
                    has_format = check_and_clean_answer_formatting(run)
                    t_node = run.find('w:t', namespaces=WORD_NS)
                    if t_node is not None and t_node.text:
                        t_node.text = t_node.text.replace('*', '').replace('∗', '')
                        t_node.text = re.sub(r'\(\s*đ(?:úng)?\s*\)', '', t_node.text, flags=re.IGNORECASE)
                        if has_format: current_opt['is_correct'] = True
            else:
                if current_opt is not None:
                    current_opt['xml'].append(el)
                    p_text = get_text_from_element(el)
                    if '*' in p_text or '∗' in p_text or re.search(r'\(\s*đ(?:úng)?\s*\)', p_text, re.IGNORECASE):
                         current_opt['is_correct'] = True
                    for run in el.findall('.//w:r', namespaces=WORD_NS):
                        if check_and_clean_answer_formatting(run): current_opt['is_correct'] = True
                        t_node = run.find('w:t', namespaces=WORD_NS)
                        if t_node is not None and t_node.text:
                            t_node.text = t_node.text.replace('*', '').replace('∗', '')
                            t_node.text = re.sub(r'\(\s*đ(?:úng)?\s*\)', '', t_node.text, flags=re.IGNORECASE)
                else: stem.append(el)
        else:
            if current_opt is not None: current_opt['xml'].append(el)
            else: stem.append(el)
                
    if current_opt is not None: options.append(current_opt)
    for opt in options:
        while len(opt['xml']) > 1 and not get_text_from_element(opt['xml'][-1]).strip(): opt['xml'].pop()

    if len(options) != 4: return block, "A", f"{zone_type} - {question_text} LỖI ĐỊNH DẠNG: Yêu cầu 4 đáp án tách rời."
    
    correct_count = sum(1 for opt in options if opt['is_correct'])
    if zone_type == "P1":
        if correct_count == 0:
            return block, "A", f"PHẦN I - {question_text} CHƯA có đáp án đúng (thiếu dấu *)."
        elif correct_count > 1:
            return block, "A", f"PHẦN I - {question_text} LỖI LOGIC: Có đến {correct_count} đáp án đúng. Phần I chỉ cho phép DUY NHẤT 1 đáp án đúng!"
            
    random.shuffle(options)
    ans_result = ""
    
    for idx, opt in enumerate(options):
        first_p = opt['xml'][0]
        label_replaced = False
        for run in first_p.findall('.//w:r', namespaces=WORD_NS):
            t_node = run.find('w:t', namespaces=WORD_NS)
            if t_node is not None and t_node.text and not label_replaced:
                sub_pattern = r'^[\s]*([A-D]|[a-d])[.)]'
                if re.search(sub_pattern, t_node.text, re.IGNORECASE):
                    separator = '.' if zone_type == "P1" else ')'
                    clean_text = re.sub(sub_pattern, '', t_node.text, count=1, flags=re.IGNORECASE).lstrip()
                    
                    new_run = copy.deepcopy(run)
                    new_t = new_run.find('w:t', namespaces=WORD_NS)
                    if new_t is not None: 
                        new_t.text = clean_text
                        new_t.set(qn('xml:space'), 'preserve')
                    remove_bold(new_run)
                    
                    t_node.text = f"{labels[idx]}{separator} "
                    t_node.set(qn('xml:space'), 'preserve') 
                    make_run_bold(run)
                    
                    if clean_text: run.addnext(new_run)
                    label_replaced = True
                    
        if zone_type == "P1":
            if opt['is_correct']: ans_result = labels[idx]
        else:
            ans_result += "Đ" if opt['is_correct'] else "S"

    has_br = False
    for opt in options:
        for el in opt['xml']:
            if el.find(f'.//{{{WORD_NS["w"]}}}br') is not None:
                has_br = True; break

    can_merge = all(len(opt['xml']) == 1 for opt in options)
    if zone_type == "P2" or not can_merge or has_br: layout = 1 
    else:
        max_len = max(len(get_text_from_element(opt['xml'][0])) for opt in options)
        has_complex = any(analyze_complexity(opt['xml'][0]) for opt in options)
        if has_complex: layout = 2 if max_len <= 20 else 1
        else:
            if max_len <= 12: layout = 4     
            elif max_len <= 40: layout = 2   
            else: layout = 1                 
            
    new_block = stem.copy()

    if layout == 1:
        for opt in options: new_block.extend(opt['xml'])
    elif layout == 2:
        table = create_invisible_table(doc, 2, 2)
        tbl_element = table._tbl
        tbl_element.getparent().remove(tbl_element)
        for idx in range(4):
            cell = table.cell(idx // 2, idx % 2)
            cell._element.remove(cell.paragraphs[0]._element)
            for el in options[idx]['xml']: 
                if el.tag.endswith('p'): clean_paragraph_for_table(el)
                cell._element.append(el)
        new_block.append(tbl_element)
    elif layout == 4:
        table = create_invisible_table(doc, 1, 4)
        tbl_element = table._tbl
        tbl_element.getparent().remove(tbl_element)
        for idx in range(4):
            cell = table.cell(0, idx)
            cell._element.remove(cell.paragraphs[0]._element)
            for el in options[idx]['xml']: 
                if el.tag.endswith('p'): clean_paragraph_for_table(el)
                cell._element.append(el)
        new_block.append(tbl_element)

    return new_block, ans_result or "A", None

def shuffle_engine(doc, parsed_data, config_data):
    ans_key, errors = [], []
    q_counter = 1
    
    for z in ["P1", "P2", "P3", "P4"]:
        
        # Chỉ xáo trộn và bóc đáp án cho P1, P2, P3
        if z in ["P1", "P2", "P3"]:
            for q_obj in parsed_data[z]:
                q_text_short = get_text_from_element(q_obj['xml'][0]).strip()[:40] + "..."
                if z in ["P1", "P2"]:
                    new_block, ans, err = process_options_and_extract_p1_p2(doc, q_obj['xml'], z, q_text_short)
                    q_obj['xml'] = new_block; q_obj['ans'] = ans
                    if err: errors.append(err)
                else:
                    new_block, ans = [], None
                    for el in q_obj['xml']:
                        is_key_line = False
                        if el.tag.endswith('p'):
                            match = re.search(r'^\s*(?:Đáp án|ĐS|Key)\s*[:=]\s*(.*)', get_text_from_element(el).strip(), re.IGNORECASE)
                            if match: ans = match.group(1).strip(); is_key_line = True
                        if not is_key_line: new_block.append(el)
                    q_obj['xml'] = new_block; q_obj['ans'] = ans or "..."
                    if not ans: errors.append(f"{z} - {q_text_short} CHƯA có dòng đáp án (Key: 123).")
                
            random.shuffle(parsed_data[z])
        
        # =========================================================================================
        # [MỚI] AI DÒ MÌN "CÂU...": CHẤP NHẬN MỌI FILE WORD BỊ PHÂN MẢNH XML CHO CẢ P4 TỰ LUẬN
        # =========================================================================================
        for index, q_dict in enumerate(parsed_data[z]):
            first_paragraph = q_dict['xml'][0] 
            p_text = get_text_from_element(first_paragraph)
            
            # Quét toàn bộ khối chữ của dòng đầu tiên để dò tìm "Câu X:"
            match = re.search(r'^(\s*)(Câu\s+\d+[:.]?)', p_text, re.IGNORECASE)
            if match:
                leading_spaces = match.group(1)
                full_match = match.group(2)
                chars_to_remove = len(leading_spaces) + len(full_match)
                
                # Cắt bỏ nhãn "Câu X:" cũ đang nằm rải rác trong nhiều mã XML
                for run in first_paragraph.findall('.//w:r', namespaces=WORD_NS):
                    t_node = run.find('w:t', namespaces=WORD_NS)
                    if t_node is not None and t_node.text:
                        if chars_to_remove > 0:
                            run_text_len = len(t_node.text)
                            if run_text_len <= chars_to_remove:
                                chars_to_remove -= run_text_len
                                t_node.text = ""
                            else:
                                t_node.text = t_node.text[chars_to_remove:].lstrip()
                                chars_to_remove = 0
                
                if config_data.get("resetChiSo", True):
                    new_label = f'{config_data.get("nhanCau", "Câu")} {index + 1}'
                else:
                    new_label = full_match.replace(':', '').replace('.', '').strip()
                separator = ':' if ':' in full_match else ('.' if '.' in full_match else ':')
                
                # Bơm lại nhãn "Câu X:" nguyên khối, in đậm và ép luôn khoảng trắng
                new_run = OxmlElement('w:r')
                rPr = OxmlElement('w:rPr')
                b = OxmlElement('w:b')
                rPr.append(b)
                new_run.append(rPr)
                
                t = OxmlElement('w:t')
                t.set(qn('xml:space'), 'preserve')
                t.text = f"{leading_spaces}{new_label}{separator} "
                new_run.append(t)
                
                pPr = first_paragraph.find(f'{{{WORD_NS["w"]}}}pPr')
                if pPr is not None:
                    pPr.addnext(new_run)
                else:
                    first_paragraph.insert(0, new_run)
        
        # Chỉ tạo key cho 3 phần đầu
        if z in ["P1", "P2", "P3"]:
            for q_obj in parsed_data[z]:
                score = "0.25" if z == "P1" else ("0.1 0.25 0.5 1" if z == "P2" else "0.5")
                ans_key.append({'q_num': q_counter, 'ans': q_obj['ans'], 'score': score, 'zone': z})
                q_counter += 1

    return parsed_data, ans_key, errors

# =====================================================================
# MODULE 5: RENDERER & GLOBAL FORMATTING (TIMES NEW ROMAN + SPACING)
# =====================================================================

def apply_global_formatting(doc):
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(1.5)
        section.footer_distance = Cm(1.27)

    def force_format(p):
        xml_str = p._element.xml
        has_complex = 'm:oMath' in xml_str or 'w:drawing' in xml_str or 'v:imagedata' in xml_str or 'w:pict' in xml_str
        
        # AI Quét để nhận biết đâu là dòng Tiêu đề (PHẦN I, PHẦN II...)
        p_text = p.text.strip().upper()
        is_header = bool(re.match(r'^(\[P[1-4]\]\s*)?PHẦN\s+(I|II|III|IV|1|2|3|4|MỘT|HAI|BA|BỐN)\b', p_text))
        
        if not has_complex:
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            
            # [CẬP NHẬT] Tiêu đề 6pt, còn lại ôm sát 0pt cực chuẩn
            if is_header:
                p.paragraph_format.space_before = Pt(6) 
                p.paragraph_format.space_after = Pt(6)
            else:
                p.paragraph_format.space_before = Pt(0) 
                p.paragraph_format.space_after = Pt(0)
            
        # Ép TẤT CẢ các chữ về TIMES NEW ROMAN
        for run in p.runs:
            run.font.name = 'Times New Roman'
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn('w:ascii'), 'Times New Roman')
            rFonts.set(qn('w:hAnsi'), 'Times New Roman')
            rFonts.set(qn('w:cs'), 'Times New Roman')
            
            # Ép về cỡ 12 (Trừ trường hợp chữ Mã đề đang là cỡ 14)
            if run.font.size != Pt(14):
                run.font.size = Pt(12)

    for p in doc.paragraphs: force_format(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs: force_format(p)

def render_template(doc, parsed_data, config_data, current_ma_de):
    body = doc._body._body
    body.clear_content()

    build_standard_header(doc, config_data, current_ma_de)

    for z in ["P1", "P2", "P3", "P4"]:
        for el in parsed_data[f"{z}_header"]: 
            text_upper = get_text_from_element(el).strip().upper()
            if "PHẦN" in text_upper:
                for run in el.findall('.//w:r', namespaces=WORD_NS):
                    make_run_bold(run)
            body.append(el)
            
        for q_obj in parsed_data[z]:
            for el in q_obj['xml']: body.append(el)

    temp_doc = Document()
    temp_doc.add_paragraph() 
    
    p_het = temp_doc.add_paragraph("---Hết---")
    p_het.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_het.runs[0].bold = True
    
    p_note1 = temp_doc.add_paragraph("- Cán bộ coi thi không giải thích gì thêm.")
    p_note1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_note1.runs[0].italic = True
    
    p_note2 = temp_doc.add_paragraph("- Học sinh không được sử dụng tài liệu.")
    p_note2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_note2.runs[0].italic = True
    
    for p in temp_doc.paragraphs:
        body.append(p._element)

    # CHUẨN HOÁ SPACING & FONT
    apply_global_formatting(doc)

    for section in doc.sections:
        header = section.header
        for p in header.paragraphs: p.text = "" 
        
        footer = section.footer
        footer.is_linked_to_previous = False
        for p in footer.paragraphs: p.text = "" 
        
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        def style_footer_run(r):
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)
            rPr = r._element.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn('w:ascii'), 'Times New Roman')
            rFonts.set(qn('w:hAnsi'), 'Times New Roman')
            rFonts.set(qn('w:cs'), 'Times New Roman')
            
        r1 = fp.add_run(f"Mã đề: {current_ma_de} - Trang ")
        style_footer_run(r1)
        
        r_page = fp.add_run()
        for e in create_field_code_element('PAGE'): r_page._r.append(e)
        style_footer_run(r_page)
        
        r2 = fp.add_run(" / ")
        style_footer_run(r2)
        
        r_numpages = fp.add_run()
        for e in create_field_code_element('NUMPAGES'): r_numpages._r.append(e)
        style_footer_run(r_numpages)

    return doc

@app.post("/api/mix-docx")
async def mix_docx_endpoint(file: UploadFile = File(...), config: str = Form(...)):
    try:
        content = await file.read()
        config_data = json.loads(config)
        so_de = int(config_data.get("soDe", 1))
        ma_de_list = config_data.get("maDeList", ["101"])
        
        if "thoiGian" not in config_data: config_data["thoiGian"] = "90"
        
        zip_buffer = io.BytesIO()
        all_exams_data = {} 
        
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
            for i in range(so_de):
                ma_de = ma_de_list[i] if i < len(ma_de_list) else str(100 + i)
                doc = Document(io.BytesIO(content))
                
                parsed_data = parse_docx(doc)
                shuffled_data, ans_key, errors = shuffle_engine(doc, parsed_data, config_data)
                
                if errors:
                    unique_errors = list(dict.fromkeys(errors))
                    return JSONResponse(status_code=400, content={"message": "Phát hiện lỗi Đề Gốc!", "details": unique_errors})
                
                final_doc = render_template(doc, shuffled_data, config_data, ma_de)
                
                all_exams_data[ma_de] = ans_key
                doc_buffer = io.BytesIO()
                final_doc.save(doc_buffer)
                zip_file.writestr(f"De_Ma_{ma_de}.docx", doc_buffer.getvalue())
            
            # ====================================================================
            # 1. FILE EXCEL: ĐÁP ÁN DỌC 
            # ====================================================================
            wb_doc = Workbook()
            ws_doc = wb_doc.active
            ws_doc.title = "Dap An Doc"
            ws_doc.append(['Mã đề', 'Câu hỏi', 'Đáp án', 'Điểm'])
            for m_de, ans_list in all_exams_data.items():
                for item in ans_list: 
                    ws_doc.append([m_de, item['q_num'], item['ans'], item['score']])
            
            doc_excel_buffer = io.BytesIO()
            wb_doc.save(doc_excel_buffer)
            doc_excel_buffer.seek(0)
            zip_file.writestr("DapAn_ChiTiet_Doc.xlsx", doc_excel_buffer.read())

            # ====================================================================
            # 2. FILE EXCEL: ĐÁP ÁN NGANG 
            # ====================================================================
            wb_ngang = Workbook()
            ws_ngang = wb_ngang.active
            ws_ngang.title = "Dap An Ngang"
            made_keys = list(all_exams_data.keys())
            
            ws_ngang.append(['Câu hỏi'] + made_keys + ['diem'])
            if len(made_keys) > 0:
                max_questions = max(len(all_exams_data[k]) for k in made_keys)
                for q_idx in range(max_questions):
                    row = [str(q_idx + 1)]
                    for m_de in made_keys: 
                        if q_idx < len(all_exams_data[m_de]): row.append(all_exams_data[m_de][q_idx]['ans'])
                        else: row.append("")
                    if q_idx < len(all_exams_data[made_keys[0]]): row.append(all_exams_data[made_keys[0]][q_idx]['score'])
                    else: row.append("")
                    ws_ngang.append(row)
            
            ngang_excel_buffer = io.BytesIO()
            wb_ngang.save(ngang_excel_buffer)
            ngang_excel_buffer.seek(0)
            zip_file.writestr("DapAn_DeTron_Ngang.xlsx", ngang_excel_buffer.read())

            # ====================================================================
            # 3. FILE EXCEL: CHUẨN OLM 
            # ====================================================================
            wb_olm = Workbook()
            ws_olm = wb_olm.active
            ws_olm.title = "Dap An OLM"
            
            if len(made_keys) > 0:
                first_made = made_keys[0]
                first_ans_list = all_exams_data[first_made]
                
                p1_list = [item for item in first_ans_list if item['zone'] == 'P1']
                p2_list = [item for item in first_ans_list if item['zone'] == 'P2']
                p3_list = [item for item in first_ans_list if item['zone'] == 'P3']
                
                num_p1 = len(p1_list)
                num_p2 = len(p2_list)
                num_p3 = len(p3_list)
                
                row1 = [""]
                if num_p1 > 0:
                    row1.extend(["Phần Ⅰ: Mỗi câu 0.25đ"] + [""] * (num_p1 - 1))
                if num_p2 > 0:
                    row1.extend(["Phần Ⅱ: Mỗi câu tối đa 1đ: đúng 1 ý 0.1đ, đúng 2 ý: 0.25đ, đúng 3 ý: 0.5đ, đúng 4 ý: 1đ."] + [""] * (num_p2 * 4 - 1))
                if num_p3 > 0:
                    row1.extend(["Phần Ⅲ: Mỗi câu 0.5 điểm"] + [""] * (num_p3 - 1))
                ws_olm.append(row1)
                
                row2 = [""]
                for i in range(1, num_p1 + 1): row2.append(str(i))
                for i in range(1, num_p2 + 1): row2.extend([f"{i}a", f"{i}b", f"{i}c", f"{i}d"])
                for i in range(1, num_p3 + 1): row2.append(f"Câu {i}")
                ws_olm.append(row2)
                
                row3 = ["Điểm"]
                for _ in range(num_p1): row3.append("0.25")
                for _ in range(num_p2 * 4): row3.append("0.25")
                for _ in range(num_p3): row3.append("0.5")
                ws_olm.append(row3)
                
                for m_de in made_keys:
                    ans_list = all_exams_data[m_de]
                    row_data = [m_de]
                    for item in ans_list:
                        if item['zone'] == 'P1':
                            row_data.append(item['ans'])
                        elif item['zone'] == 'P2':
                            ans_str = str(item['ans']).strip()
                            ans_str = (ans_str + "SSSS")[:4] 
                            for char in ans_str:
                                row_data.append(char)
                        elif item['zone'] == 'P3':
                            row_data.append(item['ans'])
                    ws_olm.append(row_data)

            olm_excel_buffer = io.BytesIO()
            wb_olm.save(olm_excel_buffer)
            olm_excel_buffer.seek(0)
            zip_file.writestr("DapAn_OLM.xlsx", olm_excel_buffer.read())

        zip_buffer.seek(0)
        return StreamingResponse(
            zip_buffer, 
            media_type="application/zip", 
            headers={'Content-Disposition': 'attachment; filename="Tap_De_Thi.zip"'}
        )

    except Exception as e:
        traceback.print_exc()
        return JSONResponse(status_code=500, content={"message": "Lỗi hệ thống", "details": [str(e)]})
